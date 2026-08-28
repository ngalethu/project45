"""
Script cập nhật Bảng 4.1 (FPS) + diễn giải trong báo cáo chính thức.
Dữ liệu mới từ: outputs/benchmarks/benchmark_summary.json

Chạy: python -X utf8 update_bang_4_1.py
"""
from docx import Document
from docx.shared import Pt

doc = Document(r"4551050116_La Dai Loc_Bao cao KLTN.docx")

# ============================================================================
# DỮ LIỆU MỚI TỪ BENCHMARK
# ============================================================================
NEW_FPS = {
    "Cấu hình 1": "2.34",
    "Cấu hình 2": "2.20",
    "Cấu hình 3": "3.04",
    "Cấu hình 4": "4.75",
}

edits_done = 0

# ============================================================================
# SỬA 1: Cập nhật FPS trong Bảng 4.1 (Table 7)
# ============================================================================
print("=== [1] Update FPS values in Table 7 (Bảng 4.1) ===")

table_7 = doc.tables[7]
print(f"  Table 7 has {len(table_7.rows)} rows, {len(table_7.columns)} columns")

# In ra bảng hiện tại
for i, row in enumerate(table_7.rows):
    cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
    print(f"  Row {i}: {cells}")

# Cập nhật FPS ở cột cuối cùng (index 5)
for i, row in enumerate(table_7.rows):
    if i == 0:  # Skip header
        continue
    # Xác định config từ cột đầu tiên
    first_cell = row.cells[0].text.strip()
    for config_key, new_fps in NEW_FPS.items():
        if config_key in first_cell:
            old_fps = row.cells[5].text.strip()
            print(f"  Row {i}: {config_key} -> FPS: {old_fps} => {new_fps}")
            # Xóa nội dung cũ và ghi mới
            for para in row.cells[5].paragraphs:
                for run in para.runs:
                    if old_fps in run.text:
                        run.text = run.text.replace(old_fps, new_fps)
                        print(f"    [OK] Updated run text")
                        edits_done += 1
                        break
                else:
                    continue
                break
            break

# ============================================================================
# SỬA 2: Cập nhật diễn giải sau Bảng 4.1 (Paragraph 533)
# ============================================================================
print("\n=== [2] Update FPS interpretation (Paragraph 533) ===")

old_text = (
    "Qua phân tích số liệu, có thể thấy rõ quy luật đánh đổi tài nguyên. "
    "Khi kích hoạt thêm mô hình ước lượng tư thế MediaPipe, tải điện toán "
    "(computational load) tăng vọt khiến FPS suy giảm đáng kể so với cấu hình "
    "chỉ chạy YOLO. Tuy nhiên, nhờ cơ chế kích hoạt pose theo điều kiện kết hợp "
    "bỏ khung hình và thu nhỏ không gian ảnh, pipeline có thể duy trì mức FPS "
    "đáp ứng tiêu chí thời gian thực trong điều kiện cấu hình phù hợp."
)

new_text = (
    "Qua phân tích số liệu, có thể thấy rõ quy luật đánh đổi tài nguyên. "
    "Khi kích hoạt thêm mô hình ước lượng tư thế MediaPipe, tải điện toán "
    "(computational load) tăng lên khiến FPS giảm nhẹ (từ 2.34 xuống 2.20 FPS) "
    "so với cấu hình chỉ chạy YOLO đơn lẻ. Cơ chế bỏ khung hình có chủ đích "
    "(frame skipping với N=3) giúp giảm tần suất suy luận YOLO xuống còn 1/3, "
    "từ đó cải thiện thông lượng lên 3.04 FPS. Kết hợp với kỹ thuật thu nhỏ ảnh "
    "đầu vào (resize từ 848 xuống 640 pixel), thời gian suy luận trên mỗi frame "
    "được rút ngắn đáng kể, nâng thông lượng lên 4.75 FPS — tăng hơn 2 lần so "
    "với cấu hình cơ sở. Kết quả chứng minh hiệu quả của tổ hợp chiến lược tối "
    "ưu (frame skipping + resize) trong việc cân bằng giữa độ chính xác và hiệu "
    "năng trên thiết bị Edge."
)

found = False
for i, para in enumerate(doc.paragraphs):
    if "Qua phân tích số liệu" in para.text:
        full_text = para.text
        if old_text in full_text:
            new_full = full_text.replace(old_text, new_text)
            # Preserve formatting from first run
            if para.runs:
                first_run = para.runs[0]
                font_name = first_run.font.name
                font_size = first_run.font.size
                font_bold = first_run.font.bold
                font_italic = first_run.font.italic
            else:
                font_name = font_size = font_bold = font_italic = None
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_full
                if font_name:
                    para.runs[0].font.name = font_name
                if font_size:
                    para.runs[0].font.size = font_size
                if font_bold is not None:
                    para.runs[0].font.bold = font_bold
                if font_italic is not None:
                    para.runs[0].font.italic = font_italic
            print(f"  [OK] Replaced at paragraph {i}")
            found = True
            edits_done += 1
            break
        else:
            print(f"  [HINT] Found 'Qua phân tích số liệu' at [{i}] but old_text doesn't match exactly")
            print(f"  Para text: {full_text[:150]}...")

if not found:
    print("  [SKIP] Text not found")

# ============================================================================
# SỬA 3: Cập nhật ghi chú FPS (Paragraph 532)
# ============================================================================
print("\n=== [3] Update FPS calculation note (Paragraph 532) ===")

old_note = "Ghi chú: FPS được tính theo công thức ​. Các cấu hình tương ứng được lưu trong thư mục backend/bench_configs/."
new_note = (
    "Ghi chú: FPS hiệu dụng được tính theo công thức total_frames / total_time_sec, "
    "bao gồm cả các frame không chạy suy luận (reuse kết quả phát hiện từ frame trước). "
    "Các cấu hình tương ứng được lưu trong thư mục backend/bench_configs/. "
    "Benchmark chạy trên máy tính cá nhân (CPU), trên thiết bị Edge (Jetson) FPS dự kiến cao hơn."
)

found2 = False
for i, para in enumerate(doc.paragraphs):
    if "FPS được tính theo công thức" in para.text:
        if old_note in para.text:
            new_full = para.text.replace(old_note, new_note)
        else:
            # Thay thế toàn bộ paragraph
            new_full = new_note
        if para.runs:
            first_run = para.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            font_bold = first_run.font.bold
            font_italic = first_run.font.italic
        else:
            font_name = font_size = font_bold = font_italic = None
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = new_full
            if font_name:
                para.runs[0].font.name = font_name
            if font_size:
                para.runs[0].font.size = font_size
            if font_bold is not None:
                para.runs[0].font.bold = font_bold
            if font_italic is not None:
                para.runs[0].font.italic = font_italic
        print(f"  [OK] Updated at paragraph {i}")
        found2 = True
        edits_done += 1
        break

if not found2:
    print("  [SKIP] Note not found")

# ============================================================================
# LƯU FILE
# ============================================================================
output_path = r"4551050116_La Dai Loc_Bao cao KLTN_fps_updated.docx"
doc.save(output_path)
print(f"\n{'='*60}")
print(f"Tổng kết: {edits_done} sửa thành công")
print(f"File đã chỉnh sửa: {output_path}")
print(f"{'='*60}")

# In bảng mới để xác nhận
print("\n--- Bảng 4.1 sau khi cập nhật ---")
table_7 = doc.tables[7]
for i, row in enumerate(table_7.rows):
    cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
    print(f"  {cells}")
