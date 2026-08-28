"""
Script chỉnh sửa lần 3:
- Sửa lỗi text (công thức toán, kích thước ảnh)
- Mở rộng mô tả DB + Dashboard
- Thêm bảng mới (Edge vs Cloud, 13 Landmarks, Config, Events)
Chạy: python -X utf8 fix_report_v3.py
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document(r"báo cáo_final.docx")

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def replace_paragraph_text(para, old_text, new_text):
    full_text = para.text
    if old_text not in full_text:
        return False
    new_full = full_text.replace(old_text, new_text)
    if para.runs:
        first_run = para.runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
        font_bold = first_run.font.bold
        font_italic = first_run.font.italic
    else:
        font_name = font_size = font_bold = font_italic = None
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_full
        if font_name:
            para.runs[0].font.name = font_name
        if font_size:
            para.runs[0].font.size = font_size
        if font_bold is not None:
            para.runs[0].font.bold = font_bold
        if font_italic is not None:
            para.runs[0].font.italic = font_italic
    return True


def set_cell_shading(cell, color):
    """Đặt màu nền cho cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_text(cell, text, bold=False, font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Đặt text cho cell với formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "Times New Roman"
    # Set spacing
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def create_table_after_paragraph(doc, para_index, headers, rows, col_widths=None, header_color="1a2744"):
    """Tạo bảng mới sau paragraph tại vị trí para_index."""
    # Get the paragraph element
    para = doc.paragraphs[para_index]
    para_element = para._element

    # Create table
    num_cols = len(headers)
    num_rows = len(rows) + 1  # +1 for header

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style the table
    table.style = doc.styles['Table Grid']

    # Set header row
    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_text(cell, header, bold=True, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, header_color)
        # White text for header
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Set data rows
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i + 1, j)
            set_cell_text(cell, str(cell_text), font_size=9)
            # Alternate row shading
            if i % 2 == 0:
                set_cell_shading(cell, "F0F4FA")

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    # Move table to after the target paragraph
    # The table was added at the end, we need to move it
    body = doc.element.body
    table_element = table._tbl

    # Find position after para_element
    body.remove(table_element)
    para_element.addnext(table_element)

    return table


def add_paragraph_after(doc, para_index, text, style='Normal', bold=False, font_size=11):
    """Thêm paragraph mới sau paragraph tại vị trí para_index."""
    para = doc.paragraphs[para_index]
    para_element = para._element

    # Create new paragraph
    new_para = doc.add_paragraph()
    new_para.style = doc.styles[style]
    run = new_para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "Times New Roman"

    # Move to after target paragraph
    body = doc.element.body
    new_para_element = new_para._element
    body.remove(new_para_element)
    para_element.addnext(new_para_element)

    return new_para


edits_done = 0

# ============================================================================
# SỬA 1: Para 85 - Công thức bounding box bị thiếu
# ============================================================================
print("=== [1] Fix bounding box formula (Para 85) ===")
para = doc.paragraphs[85]
old = "một hộp giới hạn (bounding box) được mô tả bởi bốn giá trị tọa độ là . Trong đó,  biểu diễn tọa độ tâm"
new = "một hộp giới hạn (bounding box) được mô tả bởi bốn giá trị tọa độ là (x, y, w, h). Trong đó, (x, y) biểu diễn tọa độ tâm"
if replace_paragraph_text(para, old, new):
    print("  [OK]")
    edits_done += 1
else:
    print("  [SKIP]")

# ============================================================================
# SỬA 2: Para 96 - Công thức class probabilities bị thiếu
# ============================================================================
print("\n=== [2] Fix class probabilities formula (Para 96) ===")
para = doc.paragraphs[96]
old = "mỗi ô lưới cũng dự đoán  xác suất có điều kiện"
new = "mỗi ô lưới cũng dự đoán P(Class_i | Object) xác suất có điều kiện"
if replace_paragraph_text(para, old, new):
    print("  [OK]")
    edits_done += 1
else:
    print("  [SKIP]")

# ============================================================================
# SỬA 3: Para 285 - 640x640 -> 768x768
# ============================================================================
print("\n=== [3] Fix resize 640x640 -> 768x768 (Para 285) ===")
para = doc.paragraphs[285]
if replace_paragraph_text(para, "640x640 pixel", "768x768 pixel"):
    print("  [OK]")
    edits_done += 1
else:
    print("  [SKIP]")

# ============================================================================
# SỬA 4: Mở rộng mô tả Para 267 (Database section)
# ============================================================================
print("\n=== [4] Expand DB section description (Para 267) ===")
para = doc.paragraphs[267]
old_text = "Bảng chính: Alerts"
new_text = (
    "Bảng chính: Alerts. Hệ thống sử dụng SQLite làm hệ quản trị cơ sở dữ liệu, "
    "với đường dẫn mặc định được cấu hình trong config.yaml (sqlite:///outputs/cloud_app.db). "
    "Bảng Alerts là trung tâm lưu trữ, chứa toàn bộ thông tin về các sự kiện vi phạm được phát hiện bởi phân hệ Edge và gửi lên Cloud. "
    "Mỗi bản ghi cảnh báo bao gồm metadata (loại vi phạm, thời gian, độ tin cậy, mã thiết bị), "
    "siêu dữ liệu bằng chứng (đường dẫn ảnh, video clip, file JSON), "
    "và thông tin đánh giá (trạng thái xác thực, người đánh giá, ghi chú). "
    "Chi tiết các trường dữ liệu được trình bày trong bảng bên dưới."
)
if replace_paragraph_text(para, old_text, new_text):
    print("  [OK]")
    edits_done += 1
else:
    print("  [SKIP]")

# ============================================================================
# SỬA 5: Mở rộng mô tả Para 269 (Dashboard section)
# ============================================================================
print("\n=== [5] Expand Dashboard section description (Para 269) ===")
# Para 269 is "Giao diện giám sát (Dashboard)" - it's a Heading 3
# We need to add descriptive text AFTER it
add_paragraph_after(
    doc, 269,
    "Giao diện giám sát được xây dựng theo kiến trúc Single Page Application (SPA) bằng React 19, "
    "sử dụng Vite 7 làm công cụ xây dựng. Hệ thống bao gồm năm trang chức năng chính: "
    "Dashboard (tổng quan với biểu đồ xu hướng và phân bố sự kiện), "
    "Alerts Center (quản lý cảnh báo với bộ lọc đa tiêu chí, phân trang, xác thực hàng loạt), "
    "Devices (thiết bị), Drivers (tài xế) và Settings (cài đặt). "
    "Dữ liệu được tự động làm mới mỗi 5 giây thông qua hook useDashboardData. "
    "Người quản lý có thể xem bằng chứng kỹ thuật số (ảnh/clip), "
    "thực hiện đánh giá thủ công (manual review) hoặc kích hoạt xác thực tự động bằng mô hình SlowFast. "
    "Thư viện Recharts được sử dụng để vẽ biểu đồ Bar Chart (xu hướng cảnh báo 24h) và Pie Chart (phân bố sự kiện). "
    "Lucide React cung cấp hệ thống icon thống nhất cho toàn bộ giao diện.",
    style='Normal', font_size=11
)
print("  [OK]")
edits_done += 1

# ============================================================================
# THÊM BẢNG 1: So sánh Edge AI vs Cloud AI (Chương 1, sau Para 44)
# ============================================================================
print("\n=== [6] Add Table: Edge AI vs Cloud AI comparison ===")
headers_1 = ["Tiêu chí", "Edge AI", "Cloud AI"]
rows_1 = [
    ["Độ trễ (Latency)", "< 50ms (xử lý tại chỗ)", "200-500ms+ (phụ thuộc mạng)"],
    ["Băng thông cần thiết", "Thấp (chỉ gửi metadata)", "Cao (truyền video thô)"],
    ["Bảo mật dữ liệu", "Dữ liệu tại thiết bị (riêng tư)", "Dữ liệu trên server (rủi ro)"],
    ["Hoạt động ngoại tuyến", "Có (không cần mạng)", "Không (phụ thuộc kết nối)"],
    ["Chi phí viễn thông", "Thấp", "Cao (video streaming 24/7)"],
    ["Năng lực tính toán", "Giới hạn (phần cứng nhúng)", "Gần như vô hạn (GPU cluster)"],
    ["Khả năng mở rộng", "Thêm thiết bị = thêm chi phí", "Elastic scaling (thuê cloud)"],
    ["Cập nhật mô hình", "Cần deploy từng thiết bị", "Cập nhật tập trung 1 lần"],
]
create_table_after_paragraph(doc, 44, headers_1, rows_1, col_widths=[4.5, 5.5, 5.5])
print("  [OK]")
edits_done += 1

# ============================================================================
# THÊM BẢNG 2: 13 Landmarks (Chương 2, sau Para 121)
# ============================================================================
print("\n=== [7] Add Table: 13 Landmarks details ===")
headers_2 = ["STT", "Tên điểm mốc", "Vùng cơ thể", "Vai trò trong hệ thống"]
rows_2 = [
    ["1", "nose (mũi)", "Mặt", "Xác định vị trí đầu, tính khoảng cách face proximity"],
    ["2", "left_ear (tai trái)", "Mặt", "Bổ sung ngữ cảnh vùng mặt, xác định hướng đầu"],
    ["3", "right_ear (tai phải)", "Mặt", "Bổ sung ngữ cảnh vùng mặt, xác định hướng đầu"],
    ["4", "mouth_left (miệng trái)", "Mặt", "Xác định hành vi hút thuốc (gần miệng)"],
    ["5", "mouth_right (miệng phải)", "Mặt", "Xác định hành vi hút thuốc (gần miệng)"],
    ["6", "left_shoulder (vai trái)", "Thân trên", "Xác định Driver ROI, Chest ROI, shoulder width"],
    ["7", "right_shoulder (vai phải)", "Thân trên", "Xác định Driver ROI, Chest ROI, shoulder width"],
    ["8", "left_elbow (khuỷu tay trái)", "Tay", "Ngữ cảnh thao tác tay, ước tính vị trí cổ tay"],
    ["9", "right_elbow (khuỷu tay phải)", "Tay", "Ngữ cảnh thao tác tay, ước tính vị trí cổ tay"],
    ["10", "left_wrist (cổ tay trái)", "Tay", "Tính proximity với điện thoại/thuốc lá"],
    ["11", "right_wrist (cổ tay phải)", "Tay", "Tính proximity với điện thoại/thuốc lá"],
    ["12", "left_hip (hông trái)", "Thân dưới", "Bổ sung Driver ROI, fallback shoulder width"],
    ["13", "right_hip (hông phải)", "Thân dưới", "Bổ sung Driver ROI, fallback shoulder width"],
]
create_table_after_paragraph(doc, 121, headers_2, rows_2, col_widths=[1.2, 4.0, 2.5, 7.5])
print("  [OK]")
edits_done += 1

# ============================================================================
# THÊM BẢNG 3: Tham số cấu hình mặc định (Chương 3, sau Para 224)
# ============================================================================
print("\n=== [8] Add Table: Default config parameters ===")
headers_3 = ["Tham số", "Giá trị", "Mô tả"]
rows_3 = [
    ["conf_threshold", "0.35", "Ngưỡng tin cậy tối thiểu của YOLO"],
    ["iou_threshold", "0.45", "Ngưỡng IoU cho Non-Maximum Suppression"],
    ["detect_every_n_frames", "2", "Chỉ chạy YOLO mỗi N frame (frame skipping)"],
    ["pose_every_n_frames", "3", "Chỉ chạy MediaPipe Pose mỗi N frame"],
    ["resize_width", "640", "Chiều rộng resize đầu vào trước khi suy luận"],
    ["alert_cooldown_sec", "4", "Thời gian chờ giữa 2 alert cùng loại (giây)"],
    ["phone_score_threshold", "0.62", "Ngưỡng điểm xác nhận hành vi dùng điện thoại"],
    ["smoking_score_threshold", "0.62", "Ngưỡng điểm xác nhận hành vi hút thuốc"],
    ["phone_confirm_frames", "7", "Số frame liên tiếp cần để xác nhận dùng điện thoại"],
    ["smoking_confirm_frames", "7", "Số frame liên tiếp cần để xác nhận hút thuốc"],
    ["no_seatbelt_confirm_frames", "12", "Số frame liên tiếp cần để xác nhận không thắt dây an toàn"],
    ["seatbelt_conf_threshold", "0.45", "Ngưỡng tin cậy cho phát hiện dây an toàn"],
    ["seatbelt_margin", "0.07", "Chênh lệch tối thiểu giữa seatbelt và no-seatbelt"],
    ["buffer_size", "150", "Kích thước frame buffer (deque) cho EvidenceWriter"],
    ["min_visibility", "0.35", "Ngưỡng visibility để giữ landmark từ MediaPipe"],
    ["gamma", "1.18", "Hệ số brighten gamma cho ảnh trước khi pose estimation"],
]
create_table_after_paragraph(doc, 224, headers_3, rows_3, col_widths=[5.0, 2.0, 8.0])
print("  [OK]")
edits_done += 1

# ============================================================================
# THÊM BẢNG 4: Các sự kiện vi phạm (Chương 3, sau Para 204)
# ============================================================================
print("\n=== [9] Add Table: Event types ===")
headers_4 = ["Event Type", "Mô tả", "confirm_frames", "Ngưỡng điểm", "Fallback"]
rows_4 = [
    ["using_phone", "Sử dụng điện thoại khi lái xe", "7", "0.62", "Không (cần pose)"],
    ["smoking", "Hút thuốc khi lái xe", "7", "0.62", "Có (raw conf >= 0.70)"],
    ["no_seatbelt", "Không thắt dây an toàn", "12", "0.45 + margin 0.07", "Không"],
]
create_table_after_paragraph(doc, 204, headers_4, rows_4, col_widths=[3.0, 4.5, 2.5, 2.5, 3.0])
print("  [OK]")
edits_done += 1

# ============================================================================
# THÊM MÔ TẢ NGẮN CHO HÌNH CẦN THÊM (Ghi chú cuối file)
# ============================================================================
print("\n=== [10] Add screenshot guide as comments ===")
# Add a note at the end of the document about screenshots needed
last_para = doc.paragraphs[-1]
guide_text = (
    "\n\n"
    "=== GHI CHÚ: HÌNH ẢNH CẦN CHỤP VÀ CHÈN THÊM ===\n"
    "\n"
    "Chương 2:\n"
    "- Hình: Kiến trúc mô hình YOLO (Backbone-Neck-Head) - chèn sau đoạn [97]\n"
    "- Hình: 13 landmarks trên sơ đồ cơ thể người - chèn sau bảng 13 Landmarks\n"
    "- Hình: Kiến trúc SlowFast 2-pathway (Slow + Fast + Lateral connections) - chèn sau đoạn [143]\n"
    "\n"
    "Chương 3:\n"
    "- Hình: Minh họa Driver ROI và Chest ROI trên ảnh cabin - chèn sau đoạn [215]\n"
    "\n"
    "Chương 4:\n"
    "- Screenshot: Edge pipeline đang chạy (cửa sổ OpenCV với detections + pose + FPS) - chèn sau đoạn [305]\n"
    "- Screenshot: Dashboard tổng quan (StatCards + PieChart + BarChart) - chèn sau đoạn [307]\n"
    "- Screenshot: Alerts Center (danh sách cảnh báo + bộ lọc) - chèn sau đoạn [307]\n"
    "- Screenshot: Evidence Modal (xem bằng chứng + nút Verify/Review) - chèn sau đoạn [307]\n"
    "- Ảnh: Ví dụ test scenario (phone detected, smoking detected) - chèn sau đoạn [311]\n"
    "- Ảnh: Ví dụ False Positive (nhầm điện thoại) - chèn sau đoạn [332]\n"
)

p = doc.add_paragraph()
run = p.add_run(guide_text)
run.font.size = Pt(10)
run.font.name = "Consolas"
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

print("  [OK]")
edits_done += 1

# ============================================================================
# LƯU FILE
# ============================================================================
output_path = r"báo cáo_final.docx"
doc.save(output_path)
print(f"\n{'='*60}")
print(f"Tổng kết: {edits_done} sửa/thêm thành công")
print(f"File đã chỉnh sửa: {output_path}")
print(f"{'='*60}")
