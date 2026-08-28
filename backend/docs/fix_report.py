"""
Script chỉnh sửa file báo cáo .docx để khớp với source code thực tế.
Chạy: python -X utf8 fix_report.py
"""
from docx import Document
from docx.shared import Pt

doc = Document(r"báo cáo_restored.docx")


def replace_paragraph_text(para, old_text, new_text):
    """Thay thế text trong paragraph, giữ nguyên formatting của run đầu tiên."""
    full_text = para.text
    if old_text not in full_text:
        return False

    new_full = full_text.replace(old_text, new_text)

    if para.runs:
        first_run = para.runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
        font_bold = first_run.font.bold
        font_italic = first_run.font.italic
    else:
        font_name = font_size = font_bold = font_italic = None

    for run in para.runs:
        run.text = ""

    if para.runs:
        para.runs[0].text = new_full
        if font_name:
            para.runs[0].font.name = font_name
        if font_size:
            para.runs[0].font.size = font_size
        if font_bold is not None:
            para.runs[0].font.bold = font_bold
        if font_italic is not None:
            para.runs[0].font.italic = font_italic

    return True


def replace_paragraph_preserve_runs(para, replacements):
    """Thay thế nhiều đoạn text trong paragraph."""
    full_text = para.text
    changed = False

    for old_text, new_text in replacements:
        if old_text in full_text:
            full_text = full_text.replace(old_text, new_text)
            changed = True

    if not changed:
        return False

    if para.runs:
        first_run = para.runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
        font_bold = first_run.font.bold
        font_italic = first_run.font.italic
    else:
        font_name = font_size = font_bold = font_italic = None

    for run in para.runs:
        run.text = ""

    if para.runs:
        para.runs[0].text = full_text
        if font_name:
            para.runs[0].font.name = font_name
        if font_size:
            para.runs[0].font.size = font_size
        if font_bold is not None:
            para.runs[0].font.bold = font_bold
        if font_italic is not None:
            para.runs[0].font.italic = font_italic

    return True


def find_paragraph_by_text(doc, search_text, start_idx=0):
    """Tìm paragraph chứa text cho trước, trả về (index, paragraph)."""
    for i, para in enumerate(doc.paragraphs):
        if i < start_idx:
            continue
        if search_text in para.text:
            return i, para
    return None, None


edits_done = 0
edits_failed = 0

# ============================================================================
# SỬA 1: MediaPipe 33 → 13 landmarks (Chương 2)
# ============================================================================
print("=== [1] MediaPipe landmarks: 33 -> 13 ===")
idx, para = find_paragraph_by_text(doc, "Kết quả đầu ra của MediaPipe Pose là một cấu trúc đồ thị không gian bao gồm 33 điểm mốc")
if para:
    old = "Kết quả đầu ra của MediaPipe Pose là một cấu trúc đồ thị không gian bao gồm 33 điểm mốc giải phẫu (anatomical landmarks), bao phủ toàn bộ cơ thể từ mắt, mũi, vai đến các khớp tay và chân."
    new = (
        "MediaPipe Pose cung cấp khả năng phát hiện tối đa 33 điểm mốc giải phẫu (anatomical landmarks) bao phủ toàn bộ cơ thể. "
        "Tuy nhiên, trong ngữ cảnh của hệ thống giám sát hành vi tài xế (DMS), hệ thống không khai thác toàn bộ 33 điểm mốc mà áp dụng chiến lược tinh lọc đặc trưng (feature filtering). "
        "Cụ thể, PoseEstimator chỉ trích xuất 13 điểm mốc then chốt, bao gồm: mũi (nose), tai trái/phải (left_ear, right_ear), vai trái/phải (left_shoulder, right_shoulder), "
        "khuỷu tay trái/phải (left_elbow, right_elbow), cổ tay trái/phải (left_wrist, right_wrist), hông trái/phải (left_hip, right_hip) và miệng trái/phải (mouth_left, mouth_right). "
        "Các điểm mốc này được lọc theo chỉ số visibility (ngưỡng mặc định 0.35) để loại bỏ các tọa độ bị nhiễu hoặc che khuất."
    )
    if replace_paragraph_text(para, old, new):
        print(f"  [OK] Para {idx}")
        edits_done += 1
    else:
        print(f"  [FAIL] Para {idx}")
        edits_failed += 1
else:
    print("  [SKIP]")

# ============================================================================
# SỬA 2: imgsz=640 -> 768, epochs=20 -> 100 (Chương 4)
# ============================================================================
print("\n=== [2] Tham số huấn luyện: imgsz, epochs ===")
idx, para = find_paragraph_by_text(doc, "imgsz=640")
if para:
    replacements = [
        ("imgsz=640", "imgsz=768"),
        ("epochs=20", "epochs=100"),
    ]
    if replace_paragraph_preserve_runs(para, replacements):
        print(f"  [OK] Para {idx}: imgsz=768, epochs=100")
        edits_done += 1
    else:
        print(f"  [FAIL] Para {idx}")
        edits_failed += 1
else:
    print("  [SKIP]")

# ============================================================================
# SỬA 3: Device Fleet Management (Chương 1, Cloud Node)
# Dùng index trực tiếp vì text chứa ký tự đặc biệt
# ============================================================================
print("\n=== [3] Device Fleet Management ===")
para = doc.paragraphs[68]
if "Device & Fleet Management" in para.text:
    # Lưu formatting
    if para.runs:
        first_run = para.runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
        font_bold = first_run.font.bold
        font_italic = first_run.font.italic
    else:
        font_name = font_size = font_bold = font_italic = None

    new_text = (
        "Lớp thứ ba, thuộc hướng phát triển mở rộng trong tương lai, là Cơ chế quản trị vòng đời thiết bị (Device & Fleet Management). "
        "Trong phiên bản hiện tại, phân hệ Cloud mới chỉ cung cấp điểm cuối kiểm tra trạng thái cơ bản (health check endpoint) tại đường dẫn /health, "
        "cho phép xác nhận trạng thái hoạt động của máy chủ. Các tính năng nâng cao như cơ chế theo dõi nhịp tim (Heartbeat) liên tục từ các Edge Node, "
        "giám sát mức độ tiêu hao tài nguyên (CPU/RAM), và đồng bộ hóa cấu hình từ xa (Over-The-Air / OTA updates) hiện chưa được triển khai "
        "và thuộc lộ trình phát triển tiếp theo của hệ thống."
    )

    for run in para.runs:
        run.text = ""

    if para.runs:
        para.runs[0].text = new_text
        if font_name:
            para.runs[0].font.name = font_name
        if font_size:
            para.runs[0].font.size = font_size
        if font_bold is not None:
            para.runs[0].font.bold = font_bold
        if font_italic is not None:
            para.runs[0].font.italic = font_italic

    print(f"  [OK] Para 68")
    edits_done += 1
else:
    print("  [SKIP] Text not found in para 68")

# ============================================================================
# SỬA 4: Dashboard + Risk Scoring + Push notifications (Chương 1, Cloud Node)
# ============================================================================
print("\n=== [4] Dashboard description + Risk Scoring + Push notifications ===")
idx, para = find_paragraph_by_text(doc, "Hệ thống giám sát và báo cáo trung tâm (Centralized Admin Dashboard)")
if para:
    old = (
        "Hệ thống giám sát và báo cáo trung tâm (Centralized Admin Dashboard). "
        "Nhằm chuyển hóa dữ liệu thô thành tri thức quản trị, Cloud Node cung cấp một nền tảng Web tương tác (Web-based platform) dành riêng cho các nhà điều hành đội xe (Fleet Managers). "
        "Thay vì phải tra cứu các tệp nhật ký (logs) khô khan, người quản lý được trang bị một bảng điều khiển trực quan theo thời gian thực (real-time telemetry). "
        "Hệ thống không chỉ đẩy các thông báo cảnh báo vi phạm (push notifications) ngay lập tức lên màn hình, mà còn cho phép truy xuất và xem lại các bằng chứng kỹ thuật số (digital evidence) một cách minh bạch. "
        "Sự tích lũy dữ liệu trên quy mô lớn tại tầng này tạo tiền đề cho các phân tích thống kê sâu hơn, cho phép doanh nghiệp đánh giá, xếp hạng và xây dựng hồ sơ rủi ro (Driver Risk Scoring) đối với từng cá nhân tài xế."
    )
    new = (
        "Hệ thống giám sát và báo cáo trung tâm (Centralized Admin Dashboard). "
        "Nhằm chuyển hóa dữ liệu thô thành tri thức quản trị, Cloud Node cung cấp một ứng dụng giám sát dựa trên nền tảng Web hiện đại. "
        "Về mặt kỹ thuật, giao diện giám sát được xây dựng bằng thư viện React 19 kết hợp với công cụ xây dựng Vite 8, "
        "sử dụng thư viện Recharts để vẽ biểu đồ thống kê và Lucide React cho hệ thống icon. "
        "Giao diện được thiết kế theo kiến trúc Single Page Application (SPA) với năm trang chức năng chính: "
        "Dashboard (tổng quan), Alerts Center (quản lý cảnh báo), Devices (thiết bị), Drivers (tài xế) và Settings (cài đặt). "
        "Hệ thống tự động làm mới dữ liệu mỗi 5 giây, cung cấp cái nhìn theo thời gian thực về tình hình cảnh báo. "
        "Người quản lý có thể truy xuất và xem lại các bằng chứng kỹ thuật số (digital evidence) trực tiếp từ giao diện, "
        "thực hiện xác minh thủ công (manual review) hoặc kích hoạt xác thực tự động bằng mô hình SlowFast. "
        "Chức năng xuất dữ liệu CSV và lọc cảnh báo theo nhiều tiêu chí (loại vi phạm, thiết bị, trạng thái, khoảng thời gian) cũng được hỗ trợ đầy đủ. "
        "Lưu ý: chức năng Driver Risk Scoring hiện chưa được triển khai và thuộc hướng phát triển tương lai."
    )
    if replace_paragraph_text(para, old, new):
        print(f"  [OK] Para {idx}")
        edits_done += 1
    else:
        print(f"  [FAIL] Para {idx}")
        edits_failed += 1
else:
    print("  [SKIP]")

# ============================================================================
# SỬA 5: Cloud API description (Chương 1)
# ============================================================================
print("\n=== [5] Cloud API description ===")
idx, para = find_paragraph_by_text(doc, "cung cấp hệ thống API mạnh mẽ để xây dựng các bảng điều khiển giao diện (Admin Dashboard)")
if para:
    old = "cung cấp hệ thống API mạnh mẽ để xây dựng các bảng điều khiển giao diện (Admin Dashboard), giúp nhà quản lý theo dõi và điều hành toàn bộ đội xe theo thời gian thực."
    new = "cung cấp hệ thống RESTful API (xây dựng bằng FastAPI) để phục vụ dữ liệu cho ứng dụng giám sát (Admin Dashboard) được xây dựng bằng React. Thông qua các điểm cuối API như /api/alerts, /api/statistics và /alerts/{id}/verify, hệ thống cho phép nhà quản lý truy vấn, lọc và xác thực các cảnh báo vi phạm theo thời gian thực."
    if replace_paragraph_text(para, old, new):
        print(f"  [OK] Para {idx}")
        edits_done += 1
    else:
        print(f"  [FAIL] Para {idx}")
        edits_failed += 1
else:
    print("  [SKIP]")

# ============================================================================
# SỬA 6: Phiên bản YOLO (Chương 4)
# ============================================================================
print("\n=== [6] Phiên bản YOLO ===")
idx, para = find_paragraph_by_text(doc, "mô hình YOLO của Ultralytics được sử dụng thông qua trọng số best.pt")
if para:
    old = "Ở bản triển khai hiện tại, mô hình YOLO của Ultralytics được sử dụng thông qua trọng số best.pt; phiên bản cụ thể phụ thuộc vào quá trình huấn luyện và có thể thay đổi theo lần fine-tuning."
    new = "Ở bản triển khai hiện tại, mô hình YOLO11m của Ultralytics được sử dụng thông qua trọng số best.pt (huấn luyện 100 epochs, imgsz=768, batch=16 trên Google Colab với GPU)."
    if replace_paragraph_text(para, old, new):
        print(f"  [OK] Para {idx}")
        edits_done += 1
    else:
        print(f"  [FAIL] Para {idx}")
        edits_failed += 1
else:
    print("  [SKIP]")

# ============================================================================
# SỬA 7: verify_on_cloud (Chương 3)
# ============================================================================
print("\n=== [7] verify_on_cloud config ===")
idx, para = find_paragraph_by_text(doc, "verify_on_cloud = True")
if para:
    old = "Cơ chế xác thực trên Cloud được kích hoạt khi tham số cấu hình verify_on_cloud = True và dịch vụ SlowFast khả dụng."
    new = "Cơ chế xác thực trên Cloud được kích hoạt khi người dùng chủ động gọi API verify cho từng cảnh báo (endpoint POST /alerts/{id}/verify) và dịch vụ SlowFast khả dụng. Trong file config.yaml, tham số verify_on_cloud được đặt ở chế tắt (comment) theo mặc định."
    if replace_paragraph_text(para, old, new):
        print(f"  [OK] Para {idx}")
        edits_done += 1
    else:
        print(f"  [FAIL] Para {idx}")
        edits_failed += 1
else:
    print("  [SKIP]")

# ============================================================================
# SỬA 8: SlowFast Chapter 5 (Đoạn 357)
# ============================================================================
print("\n=== [8] SlowFast trong Chương 5 ===")
idx, para = find_paragraph_by_text(doc, "cơ chế xác thực Cloud bằng SlowFast hoạt động theo heuristic mapping")
if para:
    old = "cơ chế xác thực Cloud bằng SlowFast hoạt động theo heuristic mapping từ nhãn Kinetics và chỉ kích hoạt khi cấu hình verify_on_cloud bật; việc fine-tune chuyên sâu cho hành vi tài xế nằm trong kế hoạch mở rộng."
    new = "cơ chế xác thực Cloud bằng SlowFast sử dụng heuristic mapping từ nhãn Kinetics-400 (keyword matching) để suy ra nhãn dự án. Chức năng này chỉ kích hoạt khi người dùng chủ động gọi API verify cho từng cảnh báo và dịch vụ SlowFast khả dụng. Trong phiên bản hiện tại, SlowFast chưa được fine-tune chuyên sâu cho hành vi tài xế; việc này nằm trong kế hoạch mở rộng."
    if replace_paragraph_text(para, old, new):
        print(f"  [OK] Para {idx}")
        edits_done += 1
    else:
        print(f"  [FAIL] Para {idx}")
        edits_failed += 1
else:
    print("  [SKIP]")

# ============================================================================
# SỬA 9: SlowFast thảo luận (Chương 4)
# ============================================================================
print("\n=== [9] SlowFast trong thảo luận ===")
idx, para = find_paragraph_by_text(doc, "mô hình nhận diện hành động theo thời gian (Action Recognition) như SlowFast hiện mới được thiết kế như một bằng chứng khái niệm")
if para:
    old = "mô hình nhận diện hành động theo thời gian (Action Recognition) như SlowFast hiện mới được thiết kế như một bằng chứng khái niệm (Proof of Concept) mở rộng ở tầng Cloud, chưa được tích hợp sâu vào luồng xử lý bắt buộc đối với mọi cảnh báo."
    new = "mô hình nhận diện hành động theo thời gian (Action Recognition) như SlowFast hiện mới dừng ở mức baseline (heuristic mapping từ nhãn Kinetics-400, chưa fine-tune chuyên sâu cho hành vi tài xế) và chỉ được kích hoạt theo yêu cầu tại tầng Cloud, chưa được tích hợp tự động vào luồng xử lý bắt buộc đối với mọi cảnh báo."
    if replace_paragraph_text(para, old, new):
        print(f"  [OK] Para {idx}")
        edits_done += 1
    else:
        print(f"  [FAIL] Para {idx}")
        edits_failed += 1
else:
    print("  [SKIP]")

# ============================================================================
# SỬA 10: SlowFast hạn chế (Chương 5)
# ============================================================================
print("\n=== [10] SlowFast baseline trong hạn chế ===")
idx, para = find_paragraph_by_text(doc, "Phân hệ xác thực ngữ cảnh tại Đám mây (Cloud Verification) bằng các mạng nơ-ron nhận diện chuỗi hành động")
if para:
    old = "Phân hệ xác thực ngữ cảnh tại Đám mây (Cloud Verification) bằng các mạng nơ-ron nhận diện chuỗi hành động (Action Recognition) như SlowFast hiện mới dừng ở mức baseline (heuristic mapping từ nhãn Kinetics) và chưa được fine-tune chuyên sâu cho hành vi tài xế."
    new = "Phân hệ xác thực ngữ cảnh tại Đám mây (Cloud Verification) bằng mạng nơ-ron SlowFast hiện mới dừng ở mức baseline (heuristic mapping từ nhãn Kinetics-400 thông qua keyword matching, chưa fine-tune chuyên sâu cho hành vi tài xế). Hệ thống chỉ suy ra nhãn dự án dựa trên các từ khóa liên quan trong tập 400 lớp hành động của Kinetics, do đó độ chính xác xác thực còn hạn chế đối với các hành vi đặc thù của tài xế."
    if replace_paragraph_text(para, old, new):
        print(f"  [OK] Para {idx}")
        edits_done += 1
    else:
        print(f"  [FAIL] Para {idx}")
        edits_failed += 1
else:
    print("  [SKIP]")

# ============================================================================
# SỬA 11: Bench config (Chương 4)
# ============================================================================
print("\n=== [11] Bench config ===")
idx, para = find_paragraph_by_text(doc, "FPS được tính theo hiệu dụng = total_frames/total_time_sec")
if para:
    old = "(Lưu ý: FPS được tính theo hiệu dụng = total_frames/total_time_sec. Dữ liệu đo trên video testfalse.mp4 (848x480), 4 cấu hình chạy lần lượt với detect_every_n_frames và pose_every_n_frames như bảng; cấu hình 4 dùng resize_width=640.)"
    new = "(Lưu ý: FPS được tính theo hiệu dụng = total_frames/total_time_sec. Dữ liệu đo trên video testfalse.mp4 (848x480), 4 cấu hình chạy lần lượt với detect_every_n_frames và pose_every_n_frames như bảng; cấu hình 4 dùng resize_width=640. Các cấu hình bench_config_1 đến bench_config_4 trong thư mục backend/bench_configs/ được sử dụng cho quá trình đánh giá.)"
    if replace_paragraph_text(para, old, new):
        print(f"  [OK] Para {idx}")
        edits_done += 1
    else:
        print(f"  [FAIL] Para {idx}")
        edits_failed += 1
else:
    print("  [SKIP]")

# ============================================================================
# LƯU FILE
# ============================================================================
output_path = r"báo cáo_final.docx"
doc.save(output_path)
print(f"\n{'='*60}")
print(f"Tổng kết: {edits_done} sửa thành công, {edits_failed} sửa thất bại")
print(f"File đã chỉnh sửa: {output_path}")
print(f"Backup gốc: báo cáo.docx.bak")
print(f"{'='*60}")
