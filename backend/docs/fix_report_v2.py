"""
Script chỉnh sửa lần 2: sửa lỗi Vite version và MediaPipe landmarks trong mô tả Edge libraries.
Chạy: python -X utf8 fix_report_v2.py
"""
from docx import Document

doc = Document(r"báo cáo_final.docx")


def replace_paragraph_text(para, old_text, new_text):
    full_text = para.text
    if old_text not in full_text:
        return False

    new_full = full_text.replace(old_text, new_text)

    if para.runs:
        first_run = para.runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
        font_bold = first_run.font.bold
        font_italic = first_run.font.italic
    else:
        font_name = font_size = font_bold = font_italic = None

    for run in para.runs:
        run.text = ""

    if para.runs:
        para.runs[0].text = new_full
        if font_name:
            para.runs[0].font.name = font_name
        if font_size:
            para.runs[0].font.size = font_size
        if font_bold is not None:
            para.runs[0].font.bold = font_bold
        if font_italic is not None:
            para.runs[0].font.italic = font_italic

    return True


edits_done = 0
edits_failed = 0

# ============================================================================
# SỬA 1: Vite 8 -> Vite 7 (Para 67 - Dashboard description)
# Actual: vite@7.0.6 in package.json
# ============================================================================
print("=== [1] Vite version: 8 -> 7 ===")
para = doc.paragraphs[67]
if "Vite 8" in para.text:
    if replace_paragraph_text(para, "Vite 8", "Vite 7"):
        print("  [OK] Para 67: Vite 8 -> Vite 7")
        edits_done += 1
    else:
        print("  [FAIL] Para 67")
        edits_failed += 1
else:
    print("  [SKIP] 'Vite 8' not found in para 67")

# ============================================================================
# SỬA 2: MediaPipe 33 -> 13 landmarks in Edge libraries section (Para 275)
# Para 121 already says system uses only 13, this paragraph contradicts it
# ============================================================================
print("\n=== [2] MediaPipe landmarks in Edge libraries: 33 -> 13 ===")
para = doc.paragraphs[275]
if "33 điểm mốc" in para.text:
    old = "trích xuất ma trận 33 điểm mốc giải phẫu cơ thể người lái"
    new = "trích xuất 13 điểm mốc giải phẫu then chốt của cơ thể người lái"
    if replace_paragraph_text(para, old, new):
        print("  [OK] Para 275: 33 -> 13 landmarks")
        edits_done += 1
    else:
        print("  [FAIL] Para 275")
        edits_failed += 1
else:
    print("  [SKIP] '33 điểm mốc' not found in para 275")

# ============================================================================
# LƯU FILE
# ============================================================================
output_path = r"báo cáo_final.docx"
doc.save(output_path)
print(f"\n{'='*60}")
print(f"Tổng kết: {edits_done} sửa thành công, {edits_failed} sửa thất bại")
print(f"File đã chỉnh sửa: {output_path}")
print(f"{'='*60}")
