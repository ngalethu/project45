"""
Script chỉnh sửa lần 4:
- Sửa diễn giải kết quả FPS trong mục 4.5.3 (Bảng 4.1)
- Giải thích rõ ý nghĩa của FPS khi sử dụng frame skipping
- Thêm ghi chú clarifying về cách tính FPS với frame skipping
Chạy: python -X utf8 fix_report_v4.py
"""
from docx import Document

doc = Document(r"báo cáo_final.docx")


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def replace_paragraph_text(para, old_text, new_text):
    full_text = para.text
    if old_text not in full_text:
        return False
    new_full = full_text.replace(old_text, new_text)
    if para.runs:
        first_run = para.runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
        font_bold = first_run.font.bold
        font_italic = first_run.font.italic
    else:
        font_name = font_size = font_bold = font_italic = None
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_full
        if font_name:
            para.runs[0].font.name = font_name
        if font_size:
            para.runs[0].font.size = font_size
        if font_bold is not None:
            para.runs[0].font.bold = font_bold
        if font_italic is not None:
            para.runs[0].font.italic = font_italic
    return True


edits_done = 0

# ============================================================================
# SỬA 1: Diễn giải kết quả FPS (sau Bảng 4.1)
# ============================================================================
print("=== [1] Fix FPS interpretation after Table 4.1 ===")

# Old text - the problematic interpretation that claims frame skipping
# "duy trì FPS đáp ứng tiêu chí thời gian thực" while FPS drops to 3.13
old_text_full = (
    "Qua phân tích số liệu, có thể thấy rõ quy luật đánh đổi tài nguyên. "
    "Khi kích hoạt thêm MediaPipe Pose, tải điện toán tăng vọt khiến FPS "
    "suy giảm. Tuy nhiên, nhờ frame skipping kết hợp lazy activation cho "
    "Pose, pipeline có thể duy trì FPS đáp ứng tiêu chí thời gian thực "
    "trong điều kiện cấu hình phù hợp."
)

# New text - honest interpretation that explains what FPS means with frame skipping
new_text_full = (
    "Kết quả đo lường cho thấy mức FPS giảm rõ rệt khi đồng thời kích hoạt "
    "YOLO và MediaPipe Pose (Cấu hình 2: 19.35 FPS so với Cấu hình 1: 21.46 FPS), "
    "phản ánh tải tính toán bổ sung từ mô hình Pose. Ở các cấu hình có frame skipping "
    "(N=3), chỉ số FPS đo được phản ánh thông lượng suy luận thực tế của pipeline "
    "trong điều kiện lấy mẫu thưa, không phải tốc độ hiển thị hình ảnh đầu ra. "
    "Cụ thể, mô hình YOLO chỉ được gọi suy luận mỗi 3 frame một lần, còn các frame "
    "trung gian sử dụng kết quả phát hiện từ frame trước, giúp giảm tải tài nguyên "
    "tính toán. Kỹ thuật resize ảnh đầu vào (từ 848 xuống 640 pixel) giúp rút ngắn "
    "thời gian suy luận trên mỗi frame, cải thiện thông lượng từ 3.13 lên 3.85 FPS. "
    "Hệ thống được thiết kế hướng tới triển khai trên thiết bị Edge chuyên dụng, "
    "nơi các tối ưu phần cứng (NPU, GPU nhúng) có thể tiếp tục nâng cao hiệu năng "
    "đáp ứng tiêu chí thời gian thực."
)

found = False
for i, para in enumerate(doc.paragraphs):
    if "Qua phân tích số liệu" in para.text:
        if replace_paragraph_text(para, old_text_full, new_text_full):
            print(f"  [OK] Replaced at paragraph {i}")
            found = True
            edits_done += 1
            break
        else:
            # Try with shorter text in case of formatting differences
            shorter_old = "Qua phân tích số liệu, có thể thấy rõ quy luật đánh đổi tài nguyên"
            shorter_new = (
                "Kết quả đo lường cho thấy mức FPS giảm rõ rệt khi đồng thời kích hoạt "
                "YOLO và MediaPipe Pose (Cấu hình 2: 19.35 FPS so với Cấu hình 1: 21.46 FPS), "
                "phản ánh tải tính toán bổ sung từ mô hình Pose. Ở các cấu hình có frame skipping "
                "(N=3), chỉ số FPS đo được phản ánh thông lượng suy luận thực tế của pipeline "
                "trong điều kiện lấy mẫu thưa, không phải tốc độ hiển thị hình ảnh đầu ra. "
                "Cụ thể, mô hình YOLO chỉ được gọi suy luận mỗi 3 frame một lần, còn các frame "
                "trung gian sử dụng kết quả phát hiện từ frame trước, giúp giảm tải tài nguyên "
                "tính toán. Kỹ thuật resize ảnh đầu vào (từ 848 xuống 640 pixel) giúp rút ngắn "
                "thời gian suy luận trên mỗi frame, cải thiện thông lượng từ 3.13 lên 3.85 FPS. "
                "Hệ thống được thiết kế hướng tới triển khai trên thiết bị Edge chuyên dụng, "
                "nơi các tối ưu phần cứng (NPU, GPU nhúng) có thể tiếp tục nâng cao hiệu năng "
                "đáp ứng tiêu chí thời gian thực."
            )
            if replace_paragraph_text(para, shorter_old, shorter_new):
                print(f"  [OK] Replaced (shorter match) at paragraph {i}")
                found = True
                edits_done += 1
                break

if not found:
    print("  [SKIP] Text 'Qua phân tích số liệu' not found in any paragraph")
    # Try alternative search
    for i, para in enumerate(doc.paragraphs):
        if "quy luật đánh đổi tài nguyên" in para.text:
            print(f"  [HINT] Found similar text at paragraph {i}: {para.text[:80]}...")
            break

# ============================================================================
# SỬA 2: Thêm ghi chú clarifying về cách tính FPS với frame skipping
# ============================================================================
print("\n=== [2] Add clarification to FPS calculation note ===")

old_note = "Lưu ý: FPS được tính theo hiệu dụng = total_frames/total_time_sec."
new_note = (
    "Lưu ý: FPS được tính theo hiệu dụng = total_frames/total_time_sec, "
    "bao gồm cả các frame không chạy suy luận (reuse kết quả phát hiện từ frame trước). "
    "Dữ liệu đo trên video testfalse.mp4 (848x480). Các cấu hình "
    "bench_config_1 đến bench_config_4 trong thư mục bench_configs/"
    "được sử dụng cho quá trình đánh giá."
)

found2 = False
for i, para in enumerate(doc.paragraphs):
    if "FPS được tính theo hiệu dụng" in para.text:
        if replace_paragraph_text(para, old_note, new_note):
            print(f"  [OK] Replaced at paragraph {i}")
            found2 = True
            edits_done += 1
            break

if not found2:
    print("  [SKIP] FPS note not found")

# ============================================================================
# LƯU FILE
# ============================================================================
output_path = r"báo cáo_final.docx"
doc.save(output_path)
print(f"\n{'='*60}")
print(f"Tổng kết: {edits_done} sửa/thêm thành công")
print(f"File đã chỉnh sửa: {output_path}")
print(f"{'='*60}")
